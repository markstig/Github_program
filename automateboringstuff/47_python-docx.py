# pip install python-docx

import docx

d = doc.Document('c:\\users\\al\\documents\\demo.docx')

print(d.paragraphs)

print(d.paragraphs[0].text)
print(d.paragraphs[1].text)

p = d.paragraphs[1]  # The second paragraph
print(p.runs)  # Each paragraph can have 1+ runs, if the text change the attributes of it, it will be a new 'run'. Here the attributs contains fonts, colors, italics, etc...

p.runs[0].text
p.runs[1].text  # Each run also can use 'text' to show the text out.

p.runs[1].bold
p.runs[0].bold == None
p.runs[3].italic
p.runs[3].underline = True
p.runs[3].text = 'italic and underlined.'

d.save('c:\\demo2.docx')  # Save the file.

print(p.style)  # Show the style name of the string.


import docx

d = doc.Document()  # This will create a blank new document and not on our harddrive yet.

d.add_paragraph('Hello this is a paragraph.')
d.add_paragraph('This is another paragraph.')

d.save('c:\\demo4.docx')

p = d.paragraphs[0]  # Choose the first paragraph
p.add_run('This is a new run.')
print(p.runs)

p.runs[1].bold = True
d.save('c:\\demo5.docx')


import docx

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print(getText('c:\\users\\al\\documents\\demo.docx'))
